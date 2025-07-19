#!/usr/bin/env python3
"""
Professional PDF to Word converter using pdfplumber and python-docx
Extracts text with proper formatting and structure preservation
"""

import os
import sys
import json
import pdfplumber
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import fitz  # PyMuPDF
import tempfile
import traceback

def extract_text_with_pdfplumber(pdf_path):
    """Extract text with focus on resume-like structure"""
    try:
        structured_content = []
        
        with pdfplumber.open(pdf_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                # Get simple text with layout
                text = page.extract_text(layout=True, x_tolerance=3, y_tolerance=3)
                if not text:
                    continue
                
                # Split into lines and process more simply
                lines = text.split('\n')
                
                for i, line in enumerate(lines):
                    line = line.strip()
                    if not line or len(line) < 5:
                        continue
                    
                    # Classify this line based on content patterns
                    line_type = classify_resume_line(line, i == 0)
                    
                    structured_content.append({
                        'text': line,
                        'type': line_type['type'],
                        'level': line_type.get('level'),
                        'formatting': line_type.get('formatting', {})
                    })
        
        return merge_related_content(structured_content)
    
    except Exception as e:
        print(f"pdfplumber extraction failed: {e}", file=sys.stderr)
        return None

def classify_resume_line(line, is_first_line=False):
    """Classify a line as name, section header, job title, company info, or regular text"""
    
    # Name (usually first substantial line, short, proper nouns)
    if (is_first_line or 
        (len(line.split()) <= 3 and 
         line.replace(' ', '').replace('.', '').isalpha() and
         line.istitle() and
         not any(word.lower() in ['experience', 'education', 'skills', 'summary'] for word in line.split()))):
        return {'type': 'name'}
    
    # Contact info (emails, urls, phone numbers)
    if (any(indicator in line.lower() for indicator in ['@', 'http', 'linkedin', 'behance']) or
        any(char.isdigit() for char in line) and len(line.split()) <= 8):
        return {'type': 'contact'}
    
    # Major section headers
    major_sections = ['work experience', 'experience', 'education', 'skills', 'projects', 'summary', 'objective']
    if any(section in line.lower() for section in major_sections):
        return {'type': 'section_header'}
    
    # Job titles (Designer, Developer, etc. but clean - no company mixed in)
    if (any(title in line.lower() for title in ['designer', 'developer', 'engineer', 'manager', 'analyst']) and
        len(line.split()) <= 4 and
        not any(word.lower() in ['solutions', 'technologies', 'systems', 'inc', 'ltd'] for word in line.split()) and
        not any(char.isdigit() for char in line)):
        return {'type': 'job_title'}
    
    # Company/date lines (contain years and company indicators)
    if (any(str(year) in line for year in range(2015, 2025)) and
        any(indicator in line.lower() for indicator in ['solutions', 'technologies', 'present', 'remote', 'bangladesh', 'lahore'])):
        return {'type': 'company_date'}
    
    # List items
    if line.lstrip().startswith(('•', '-', '*', '▪', '◦')):
        return {'type': 'list_item'}
    
    # Everything else is regular paragraph text
    return {'type': 'paragraph'}

def merge_related_content(content_list):
    """Merge related content that should be combined"""
    merged = []
    
    i = 0
    while i < len(content_list):
        current = content_list[i]
        
        # Merge consecutive contact info lines
        if current['type'] == 'contact':
            contact_lines = [current['text']]
            j = i + 1
            while j < len(content_list) and content_list[j]['type'] == 'contact':
                contact_lines.append(content_list[j]['text'])
                j += 1
            
            merged.append({
                'text': ' '.join(contact_lines),
                'type': 'contact',
                'formatting': {}
            })
            i = j
            continue
        
        # Merge consecutive paragraph lines that aren't job descriptions
        if current['type'] == 'paragraph':
            para_lines = [current['text']]
            j = i + 1
            # Only merge if next lines are also paragraphs and not too long combined
            while (j < len(content_list) and 
                   content_list[j]['type'] == 'paragraph' and
                   len(' '.join(para_lines + [content_list[j]['text']])) < 500):
                para_lines.append(content_list[j]['text'])
                j += 1
            
            merged.append({
                'text': ' '.join(para_lines),
                'type': 'paragraph',
                'formatting': {}
            })
            i = j
            continue
        
        # Keep other types as-is
        merged.append(current)
        i += 1
    
    return merged

def process_table_data(table):
    """Process extracted table data into structured format"""
    if not table or len(table) < 2:
        return None
    
    # Clean table data
    cleaned_table = []
    for row in table:
        if row and any(cell and str(cell).strip() for cell in row):
            cleaned_row = [str(cell).strip() if cell else "" for cell in row]
            cleaned_table.append(cleaned_row)
    
    if len(cleaned_table) < 2:
        return None
    
    return {
        'text': '',  # Will be filled during Word creation
        'type': 'table',
        'data': cleaned_table,
        'formatting': {'table_type': 'data'}
    }

def analyze_paragraph_with_structure(paragraph_lines, text):
    """Analyze paragraph with enhanced structure detection"""
    if not paragraph_lines:
        return {'text': text, 'type': 'paragraph', 'formatting': {}}
    
    # Check for list items
    is_list_item = detect_list_item(text)
    
    # Check for numbered items
    is_numbered_item = detect_numbered_item(text)
    
    # Analyze formatting across the paragraph
    font_sizes = [line['font_size'] for line in paragraph_lines]
    is_bold = any(line['is_bold'] for line in paragraph_lines)
    avg_font_size = sum(font_sizes) / len(font_sizes)
    
    # Determine paragraph type
    is_heading = detect_advanced_heading(text, paragraph_lines)
    
    # Determine the content type
    if is_heading:
        content_type = 'heading'
    elif is_numbered_item:
        content_type = 'numbered_item'
    elif is_list_item:
        content_type = 'list_item'
    else:
        content_type = 'paragraph'
    
    return {
        'text': text,
        'type': content_type,
        'level': get_heading_level(text) if is_heading else None,
        'formatting': {
            'is_bold': is_bold,
            'font_size': avg_font_size,
            'line_count': len(paragraph_lines),
            'is_numbered': is_numbered_item,
            'is_list': is_list_item
        }
    }

def detect_list_item(text):
    """Detect if text is a list item"""
    text_stripped = text.strip()
    return (
        text_stripped.startswith('•') or
        text_stripped.startswith('-') or
        text_stripped.startswith('*') or
        (len(text_stripped) > 0 and text_stripped[0] in '•▪▫◦‣⁃')
    )

def detect_numbered_item(text):
    """Detect if text starts with a number or numbered list"""
    import re
    text_stripped = text.strip()
    
    # Pattern for numbered items: "1.", "1)", "(1)", "1.1", etc.
    numbered_patterns = [
        r'^\d+\.',  # 1.
        r'^\d+\)',  # 1)
        r'^\(\d+\)',  # (1)
        r'^\d+\.\d+',  # 1.1
        r'^[a-zA-Z]\.',  # a., A.
        r'^[a-zA-Z]\)',  # a), A)
    ]
    
    return any(re.match(pattern, text_stripped) for pattern in numbered_patterns)

def should_start_new_paragraph_enhanced(line_data, current_paragraph):
    """Enhanced paragraph break detection matching iLovePDF structure"""
    if not current_paragraph:
        return False
    
    last_line = current_paragraph[-1]
    text = line_data['text']
    last_text = last_line['text']
    
    # Always break for clear list items
    if detect_list_item(text):
        return True
    
    # Break for job titles (clean, short professional titles)
    if (any(title in text.lower() for title in ['ui/ux designer', 'designer', 'developer']) and
        len(text.split()) <= 4 and not any(char.isdigit() for char in text)):
        return True
    
    # Break for company/date lines
    if (any(str(year) in text for year in range(2020, 2025)) and
        any(word in text.lower() for word in ['solutions', 'technologies', 'present', 'remote'])):
        return True
    
    # Break for major section headers
    if any(section in text.lower() for section in ['work experience', 'education', 'skills']):
        return True
    
    # Break for significant formatting changes
    font_change = abs(line_data['font_size'] - last_line['font_size']) > 2
    bold_change = line_data['is_bold'] != last_line['is_bold']
    
    return font_change or bold_change

def extract_formatted_lines(chars):
    """Extract lines with formatting information from character data"""
    if not chars:
        return []
    
    # Group characters by line (y-coordinate)
    lines = {}
    for char in chars:
        y = round(char['y0'], 1)  # Round to avoid floating point issues
        if y not in lines:
            lines[y] = []
        lines[y].append(char)
    
    # Sort lines by y-coordinate (top to bottom)
    sorted_lines = []
    for y in sorted(lines.keys(), reverse=True):  # Reverse because y increases downward in PDFs
        line_chars = sorted(lines[y], key=lambda c: c['x0'])  # Sort characters left to right
        
        # Combine characters into text with formatting info
        line_text = ''.join([c['text'] for c in line_chars])
        
        # Analyze formatting
        font_sizes = [c.get('size', 12) for c in line_chars if c.get('size')]
        font_names = [c.get('fontname', '') for c in line_chars if c.get('fontname')]
        
        avg_font_size = sum(font_sizes) / len(font_sizes) if font_sizes else 12
        is_bold = any('bold' in str(name).lower() for name in font_names)
        
        sorted_lines.append({
            'text': line_text.strip(),
            'font_size': avg_font_size,
            'is_bold': is_bold,
            'y_position': y,
            'char_count': len(line_chars)
        })
    
    return sorted_lines

def should_start_new_paragraph(line_data, current_paragraph):
    """More conservative paragraph break detection"""
    if not current_paragraph:
        return False
    
    last_line = current_paragraph[-1]
    
    # Only break on significant formatting changes
    significant_font_size_change = abs(line_data['font_size'] - last_line['font_size']) > 2
    bold_change = line_data['is_bold'] != last_line['is_bold']
    
    # Check for content patterns
    text = line_data['text']
    last_text = last_line['text']
    
    # Clear sentence ending followed by new topic
    sentence_end = last_text.rstrip().endswith(('.', '!', '?'))
    starts_capital = text and text[0].isupper()
    
    # Only break for clear new sections
    is_clear_new_section = (
        any(section in text.lower() for section in ['work experience', 'education', 'skills', 'projects']) or
        (contains_date_pattern(text) and any(word in text.lower() for word in ['designer', 'developer', 'engineer'])) or
        (text.endswith(':') and len(text.split()) <= 5)
    )
    
    # Be more conservative - only break on clear indicators
    return (significant_font_size_change or 
            (bold_change and is_clear_new_section) or
            (sentence_end and starts_capital and len(current_paragraph) > 2))

def analyze_paragraph(paragraph_lines, text):
    """Analyze a paragraph to determine its type and formatting"""
    if not paragraph_lines:
        return {'text': text, 'type': 'paragraph', 'formatting': {}}
    
    # Analyze formatting across the paragraph
    font_sizes = [line['font_size'] for line in paragraph_lines]
    is_bold = any(line['is_bold'] for line in paragraph_lines)
    avg_font_size = sum(font_sizes) / len(font_sizes)
    
    # Determine paragraph type
    is_heading = detect_advanced_heading(text, paragraph_lines)
    
    return {
        'text': text,
        'type': 'heading' if is_heading else 'paragraph',
        'level': get_heading_level(text) if is_heading else None,
        'formatting': {
            'is_bold': is_bold,
            'font_size': avg_font_size,
            'line_count': len(paragraph_lines)
        }
    }

def detect_advanced_heading(text, lines):
    """Detect headings to match iLovePDF's professional structure"""
    if not text or len(text) > 100:
        return False
    
    # Get formatting info
    avg_font_size = sum(line['font_size'] for line in lines) / len(lines)
    is_bold = any(line['is_bold'] for line in lines)
    
    # Main resume sections - these should be headings
    major_sections = ['work experience', 'experience', 'education', 'skills', 'projects', 'summary']
    is_major_section = any(section in text.lower() for section in major_sections)
    
    # Name - typically at the top, bold, larger font
    is_name = (
        len(text.split()) <= 3 and 
        text.replace(' ', '').replace('.', '').isalpha() and  # Only letters and spaces/dots
        not any(word.lower() in ['and', 'the', 'with', 'for'] for word in text.split()) and
        (is_bold or avg_font_size > 14)
    )
    
    # Job titles - but NOT with companies or dates mixed in
    is_clean_job_title = (
        any(title in text.lower() for title in ['ui/ux designer', 'designer', 'developer', 'engineer']) and
        len(text.split()) <= 4 and  # Keep it short and clean
        not any(char.isdigit() for char in text) and  # No dates mixed in
        not any(word.lower() in ['solutions', 'technologies', 'systems'] for word in text.split())
    )
    
    # Only these specific types should be headings
    definite_heading = is_major_section or is_name or is_clean_job_title
    
    # Must have formatting distinction 
    has_formatting = is_bold or avg_font_size > 12.5
    
    return definite_heading and has_formatting

def detect_heading(text, lines):
    """Detect if text is likely a heading based on various criteria"""
    if not text or len(text) > 200:  # Very long text is unlikely to be a heading
        return False
    
    # Check for common heading patterns
    heading_indicators = [
        text.isupper(),  # ALL CAPS
        len(text.split()) <= 8,  # Short phrases
        any(word in text.lower() for word in ['experience', 'education', 'skills', 'projects', 'work', 'employment']),
        text.endswith(':'),  # Ends with colon
        any(char.isdigit() for char in text) and ('20' in text),  # Contains years
    ]
    
    return sum(heading_indicators) >= 2

def get_heading_level(text):
    """Determine heading level based on content"""
    if any(word in text.lower() for word in ['experience', 'education', 'skills']):
        return 1
    elif any(char.isdigit() for char in text) and ('20' in text):
        return 2
    else:
        return 2

def is_likely_heading_or_title(line):
    """Check if line looks like a job title or section header"""
    title_patterns = [
        'designer', 'developer', 'engineer', 'manager', 'analyst', 'specialist',
        'experience', 'education', 'skills', 'projects', 'work history'
    ]
    return any(pattern in line.lower() for pattern in title_patterns)

def contains_date_pattern(line):
    """Check if line contains date patterns"""
    import re
    date_patterns = [
        r'\b(19|20)\d{2}\b',  # Years
        r'\b(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\b',  # Months
        r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b'  # Date formats
    ]
    return any(re.search(pattern, line, re.IGNORECASE) for pattern in date_patterns)

def is_section_break(current_line, previous_line):
    """Determine if there should be a section break"""
    if len(current_line) < len(previous_line) * 0.5:  # Much shorter line
        return True
    if current_line.isupper() and len(current_line) < 50:
        return True
    return False

def extract_text_with_pymupdf(pdf_path):
    """Advanced fallback extraction using PyMuPDF with better block detection"""
    try:
        paragraphs = []
        doc = fitz.open(pdf_path)
        
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            
            # Get text blocks with positioning - this preserves layout better
            blocks = page.get_text("dict")["blocks"]
            
            for block in blocks:
                if "lines" in block and block.get("type") == 0:  # Text blocks only
                    block_lines = []
                    
                    for line in block["lines"]:
                        line_text = ""
                        for span in line["spans"]:
                            if span.get("text", "").strip():
                                line_text += span["text"]
                        
                        if line_text.strip():
                            block_lines.append(line_text.strip())
                    
                    if block_lines:
                        # Each block is typically a paragraph in PyMuPDF
                        paragraph_text = ' '.join(block_lines)
                        if len(paragraph_text.strip()) > 20:  # Only substantial paragraphs
                            paragraphs.append(paragraph_text)
        
        doc.close()
        
        # Remove duplicates while preserving order
        seen = set()
        unique_paragraphs = []
        for para in paragraphs:
            if para not in seen:
                seen.add(para)
                unique_paragraphs.append(para)
        
        return unique_paragraphs
    
    except Exception as e:
        print(f"PyMuPDF extraction failed: {e}", file=sys.stderr)
        return None

def create_word_document(structured_content, output_path):
    """Create Word document matching iLovePDF's exact structure"""
    try:
        doc = Document()
        
        # Set margins for professional resume
        section = doc.sections[0]
        section.page_height = Inches(11)
        section.page_width = Inches(8.5)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        
        # Configure document styles
        configure_document_styles(doc)
        
        # Process content with proper iLovePDF-style formatting
        for item in structured_content:
            text = item['text'].strip()
            if not text or len(text) < 3:
                continue
            
            content_type = item['type']
            
            if content_type == 'name':
                create_resume_name(doc, text)
            elif content_type == 'contact':
                create_resume_contact(doc, text)
            elif content_type == 'section_header':
                create_resume_section(doc, text)
            elif content_type == 'job_title':
                create_resume_job_title(doc, text)
            elif content_type == 'company_date':
                create_resume_company_date(doc, text)
            elif content_type == 'list_item':
                create_resume_bullet(doc, text)
            elif content_type == 'table':
                if 'data' in item:
                    create_word_table(doc, item['data'])
            else:  # paragraph
                create_resume_paragraph(doc, text)
        
        # Ensure document has content
        if len(doc.paragraphs) == 0:
            doc.add_paragraph("No content could be extracted from the PDF.")
        
        doc.save(output_path)
        
        # Log structure
        types_count = {}
        for item in structured_content:
            t = item['type']
            types_count[t] = types_count.get(t, 0) + 1
        
        print(f"Created resume with: {types_count}", file=sys.stderr)
        return True
        
    except Exception as e:
        print(f"Word document creation failed: {e}", file=sys.stderr)
        print(traceback.format_exc(), file=sys.stderr)
        return False

def create_resume_name(doc, text):
    """Create name heading - large, centered, bold"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(18)
    run.bold = True
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(6)

def create_resume_contact(doc, text):
    """Create contact info - small, centered"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(12)

def create_resume_section(doc, text):
    """Create section header - left-aligned, bold"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(14)
    run.bold = True
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(18)
    para.paragraph_format.space_after = Pt(6)

def create_resume_job_title(doc, text):
    """Create job title - bold, left-aligned"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.bold = True
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(2)

def create_resume_company_date(doc, text):
    """Create company/date info - italic, right-aligned"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.italic = True
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    para.paragraph_format.space_after = Pt(4)

def create_resume_paragraph(doc, text):
    """Create regular paragraph text"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.15

def create_resume_bullet(doc, text):
    """Create bullet point"""
    clean_text = text.lstrip('•▪▫◦‣⁃-* ').strip()
    para = doc.add_paragraph(clean_text, style='List Bullet')
    para.paragraph_format.left_indent = Inches(0.25)
    para.paragraph_format.space_after = Pt(3)
    
    for run in para.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(11)

def create_word_table(doc, table_data):
    """Create a properly formatted table in Word"""
    if not table_data or len(table_data) < 1:
        return
    
    # Create table with appropriate dimensions
    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
    table.style = 'Table Grid'
    
    # Fill table data
    for row_idx, row_data in enumerate(table_data):
        for col_idx, cell_data in enumerate(row_data):
            if col_idx < len(table.rows[row_idx].cells):
                cell = table.rows[row_idx].cells[col_idx]
                cell.text = str(cell_data) if cell_data else ""
                
                # Format header row
                if row_idx == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
    
    # Add spacing after table
    doc.add_paragraph()

def create_numbered_item(doc, text, formatting):
    """Create a numbered list item"""
    para = doc.add_paragraph(text)
    para.style = 'List Number'
    
    # Apply formatting
    for run in para.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        if formatting.get('is_bold', False):
            run.bold = True

def create_list_item(doc, text, formatting):
    """Create bullet points that match iLovePDF's professional style"""
    # Remove bullet character if it exists
    clean_text = text.lstrip('•▪▫◦‣⁃-* ').strip()
    
    para = doc.add_paragraph(clean_text)
    para.style = 'List Bullet'
    para.paragraph_format.left_indent = Inches(0.25)  # Professional indent
    para.paragraph_format.space_after = Pt(3)  # Tight spacing
    
    # Apply clean formatting
    for run in para.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(11)

def configure_document_styles(doc):
    """Configure document styles for professional appearance"""
    from docx.enum.style import WD_STYLE_TYPE
    
    # Normal style
    normal_style = doc.styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Calibri'  # More modern font
    normal_font.size = Pt(11)
    
    normal_paragraph_format = normal_style.paragraph_format
    normal_paragraph_format.space_after = Pt(6)
    normal_paragraph_format.line_spacing = 1.15
    normal_paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

def create_formatted_heading(doc, text, level, formatting):
    """Create headings that match iLovePDF's clean professional style"""
    
    # Name formatting (largest, centered)
    if (len(text.split()) <= 3 and 
        text.replace(' ', '').replace('.', '').isalpha() and
        not any(word.lower() in ['experience', 'education', 'skills'] for word in text.split())):
        para = doc.add_heading(text, level=1)
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(6)
        
        for run in para.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(20)
            run.bold = True
            
    # Major section headings (Work Experience, Education, etc.)
    elif any(section in text.lower() for section in ['work experience', 'experience', 'education', 'skills', 'summary']):
        para = doc.add_heading(text, level=2)
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.space_before = Pt(18)
        para.paragraph_format.space_after = Pt(6)
        
        for run in para.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(14)
            run.bold = True
            
    # Job titles
    else:
        para = doc.add_heading(text, level=3)
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(4)
        
        for run in para.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(12)
            run.bold = True

def create_formatted_paragraph(doc, text, formatting):
    """Create paragraphs that match iLovePDF's clean style"""
    
    # Check if this is contact info (should be centered under name)
    if (any(indicator in text for indicator in ['@gmail.', 'linkedin.', 'behance.', 'http']) or
        any(char.isdigit() for char in text) and len(text) < 100):
        para = doc.add_paragraph(text)
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_after = Pt(4)
        
        for run in para.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
    
    # Company and date info (should be on separate line, right-aligned)
    elif (any(str(year) in text for year in range(2020, 2025)) and 
          any(word in text.lower() for word in ['solutions', 'technologies', 'ltd', 'inc', 'present', 'remote'])):
        para = doc.add_paragraph(text)
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        para.paragraph_format.space_after = Pt(6)
        
        for run in para.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            run.italic = True
    
    # Regular paragraph text
    else:
        para = doc.add_paragraph(text)
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.line_spacing = 1.15
        
        for run in para.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(11)



def convert_pdf_to_word(input_path, output_path):
    """Main conversion function with enhanced structure preservation"""
    try:
        # Check if file exists and is readable
        if not os.path.exists(input_path):
            return {"success": False, "error": "PDF file not found"}
        
        # Try pdfplumber first (better for structured text extraction)
        structured_content = extract_text_with_pdfplumber(input_path)
        
        if not structured_content:
            # Fallback to PyMuPDF - convert to structured format
            paragraphs = extract_text_with_pymupdf(input_path)
            if paragraphs:
                structured_content = []
                for para in paragraphs:
                    structured_content.append({
                        'text': para,
                        'type': 'heading' if detect_heading(para, [para]) else 'paragraph',
                        'level': 2
                    })
        
        if not structured_content:
            return {"success": False, "error": "No text found in PDF. The PDF may be image-based or encrypted."}
        
        # Filter out very short content
        meaningful_content = []
        for item in structured_content:
            if len(item['text'].strip()) > 8:  # Include substantial content
                meaningful_content.append(item)
        
        if not meaningful_content:
            return {"success": False, "error": "No substantial text content found in PDF"}
        
        # Create Word document with structure
        if create_word_document(meaningful_content, output_path):
            headings_count = len([item for item in meaningful_content if item['type'] == 'heading'])
            paragraphs_count = len([item for item in meaningful_content if item['type'] == 'paragraph'])
            return {
                "success": True, 
                "message": f"Successfully converted PDF to Word. Created {headings_count} headings and {paragraphs_count} paragraphs."
            }
        else:
            return {"success": False, "error": "Failed to create Word document"}
    
    except Exception as e:
        error_msg = f"Conversion failed: {str(e)}"
        print(f"Error in convert_pdf_to_word: {error_msg}", file=sys.stderr)
        print(traceback.format_exc(), file=sys.stderr)
        return {"success": False, "error": error_msg}

def main():
    """Command line interface"""
    if len(sys.argv) != 3:
        print(json.dumps({"success": False, "error": "Usage: python pdf_converter.py <input.pdf> <output.docx>"}))
        sys.exit(1)
    
    input_path = sys.argv[1]
    output_path = sys.argv[2]
    
    result = convert_pdf_to_word(input_path, output_path)
    print(json.dumps(result))

if __name__ == "__main__":
    main()
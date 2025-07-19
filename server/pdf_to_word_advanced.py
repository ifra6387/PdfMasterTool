#!/usr/bin/env python3
"""
Advanced PDF to Word converter with structure preservation
Uses pdf2docx for professional-grade conversion matching iLovePDF quality
"""

import sys
import json
import os
from pathlib import Path

def convert_pdf_to_word_professional(input_path, output_path):
    """
    Convert PDF to Word with professional structure preservation
    Uses pdf2docx library for layout-aware conversion
    """
    try:
        from pdf2docx import Converter
        
        # Create converter instance
        cv = Converter(input_path)
        
        # Convert with advanced options for better structure preservation
        cv.convert(
            output_path,
            start=0,          # Start from first page
            end=None,         # Convert all pages
            pages=None,       # Convert all pages
            # Advanced formatting options
            layout_recognition=True,    # Enable layout recognition
            table_recognition=True,     # Detect and preserve tables
            header_footer=True,         # Preserve headers and footers
            multi_processing=True,      # Use multi-processing for speed
            cpu_count=None,            # Use all available CPUs
        )
        
        # Close converter
        cv.close()
        
        print(json.dumps({
            "success": True,
            "message": "PDF converted to Word with professional structure preservation"
        }))
        return True
        
    except ImportError:
        # Fallback to enhanced pdfplumber approach if pdf2docx not available
        return convert_with_enhanced_pdfplumber(input_path, output_path)
    
    except Exception as e:
        print(f"pdf2docx conversion failed: {e}", file=sys.stderr)
        # Fallback to enhanced approach
        return convert_with_enhanced_pdfplumber(input_path, output_path)

def convert_with_enhanced_pdfplumber(input_path, output_path):
    """
    Enhanced fallback conversion using pdfplumber with better structure detection
    """
    try:
        import pdfplumber
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        import re
        
        doc = Document()
        
        # Set document margins for professional appearance
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        with pdfplumber.open(input_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                # Extract tables first
                tables = page.extract_tables()
                if tables:
                    for table_data in tables:
                        create_word_table(doc, table_data)
                
                # Extract text with layout preservation
                text = page.extract_text(layout=True, x_tolerance=3, y_tolerance=3)
                if text:
                    process_text_with_structure(doc, text, page_num == 0)
                
                # Add page break except for last page
                if page_num < len(pdf.pages) - 1:
                    doc.add_page_break()
        
        # Save document
        doc.save(output_path)
        
        print(json.dumps({
            "success": True,
            "message": "PDF converted using enhanced structure preservation"
        }))
        return True
        
    except Exception as e:
        print(f"Enhanced conversion failed: {e}", file=sys.stderr)
        print(json.dumps({
            "success": False,
            "error": f"Conversion failed: {str(e)}"
        }))
        return False

def create_word_table(doc, table_data):
    """Create properly formatted Word table"""
    if not table_data or len(table_data) < 1:
        return
    
    # Filter out empty rows
    filtered_data = []
    for row in table_data:
        if row and any(cell and str(cell).strip() for cell in row):
            cleaned_row = [str(cell).strip() if cell else "" for cell in row]
            filtered_data.append(cleaned_row)
    
    if len(filtered_data) < 1:
        return
    
    # Create table
    max_cols = max(len(row) for row in filtered_data)
    table = doc.add_table(rows=len(filtered_data), cols=max_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    
    # Fill table data
    for row_idx, row_data in enumerate(filtered_data):
        for col_idx, cell_data in enumerate(row_data):
            if col_idx < len(table.rows[row_idx].cells):
                cell = table.rows[row_idx].cells[col_idx]
                cell.text = str(cell_data) if cell_data else ""
                
                # Format first row as header
                if row_idx == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
    
    # Add spacing after table
    doc.add_paragraph()

def process_text_with_structure(doc, text, is_first_page=False):
    """Process text while preserving structure and formatting"""
    lines = text.split('\n')
    
    for i, line in enumerate(lines):
        line = line.strip()
        if not line:
            continue
        
        # Determine line type and formatting
        line_type = classify_line_type(line, i == 0 and is_first_page)
        
        if line_type == 'name':
            create_name_heading(doc, line)
        elif line_type == 'contact':
            create_contact_info(doc, line)
        elif line_type == 'section_header':
            create_section_header(doc, line)
        elif line_type == 'job_title_with_date':
            create_job_title_with_date(doc, line)
        elif line_type == 'bullet_point':
            create_bullet_point(doc, line)
        elif line_type == 'numbered_item':
            create_numbered_item(doc, line)
        else:
            create_paragraph(doc, line)

def classify_line_type(line, is_first_line=False):
    """Classify line type for appropriate formatting"""
    
    # Name - first line or short proper noun phrase
    if (is_first_line or 
        (len(line.split()) <= 3 and line.istitle() and 
         not any(word.lower() in ['experience', 'education', 'skills'] for word in line.split()))):
        return 'name'
    
    # Contact information
    if (any(indicator in line.lower() for indicator in ['@', 'http', 'linkedin', 'phone', 'email']) or
        re.search(r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b', line)):
        return 'contact'
    
    # Section headers
    section_patterns = ['work experience', 'experience', 'education', 'skills', 'projects', 'summary', 'objective']
    if any(pattern in line.lower() for pattern in section_patterns):
        return 'section_header'
    
    # Job title with date (contains both job title and date)
    if (any(title in line.lower() for title in ['designer', 'developer', 'engineer', 'manager']) and
        re.search(r'\b(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s*\d{4}|Present|\d{4}\s*-\s*\d{4}', line)):
        return 'job_title_with_date'
    
    # Bullet points
    if re.match(r'^\s*[•▪▫◦‣⁃-]\s*', line):
        return 'bullet_point'
    
    # Numbered items
    if re.match(r'^\s*\d+[\.)]\s*', line):
        return 'numbered_item'
    
    return 'paragraph'

def create_name_heading(doc, text):
    """Create centered name heading"""
    para = doc.add_heading(text, level=1)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in para.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(18)

def create_contact_info(doc, text):
    """Create centered contact information"""
    para = doc.add_paragraph(text)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in para.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(10)

def create_section_header(doc, text):
    """Create section header"""
    para = doc.add_heading(text, level=2)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in para.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(14)

def create_job_title_with_date(doc, text):
    """Create job title with right-aligned date"""
    # Try to separate job title and date
    date_pattern = r'(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s*\d{4}.*?Present|Present|\d{4}\s*-\s*\d{4}'
    date_match = re.search(date_pattern, text)
    
    if date_match:
        date_part = date_match.group(0).strip()
        job_title = text.replace(date_part, '').strip()
        
        # Create table for alignment
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Remove table borders
        for row in table.rows:
            for cell in row.cells:
                cell._element.get_or_add_tcPr().append(
                    cell._element._new_tcBorders()
                )
        
        # Job title (left)
        left_cell = table.cell(0, 0)
        left_cell.text = job_title
        left_para = left_cell.paragraphs[0]
        for run in left_para.runs:
            run.bold = True
            run.font.name = 'Calibri'
            run.font.size = Pt(12)
        
        # Date (right)
        right_cell = table.cell(0, 1)
        right_cell.text = date_part
        right_para = right_cell.paragraphs[0]
        right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in right_para.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
    else:
        # Fallback to regular heading
        para = doc.add_heading(text, level=3)
        for run in para.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(12)

def create_bullet_point(doc, text):
    """Create bullet point"""
    clean_text = re.sub(r'^\s*[•▪▫◦‣⁃-]\s*', '', text)
    para = doc.add_paragraph(clean_text, style='List Bullet')
    for run in para.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(11)

def create_numbered_item(doc, text):
    """Create numbered list item"""
    clean_text = re.sub(r'^\s*\d+[\.)]\s*', '', text)
    para = doc.add_paragraph(clean_text, style='List Number')
    for run in para.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(11)

def create_paragraph(doc, text):
    """Create regular paragraph"""
    para = doc.add_paragraph(text)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in para.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(11)

if __name__ == "__main__":
    if len(sys.argv) != 3:
        print(json.dumps({
            "success": False,
            "error": "Usage: python pdf_to_word_advanced.py <input_pdf> <output_docx>"
        }))
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2]
    
    if not os.path.exists(input_file):
        print(json.dumps({
            "success": False,
            "error": f"Input file not found: {input_file}"
        }))
        sys.exit(1)
    
    success = convert_pdf_to_word_professional(input_file, output_file)
    sys.exit(0 if success else 1)
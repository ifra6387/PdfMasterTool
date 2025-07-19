#!/usr/bin/env python3
"""
Professional Word to PDF Converter for Replit
Converts .docx files to PDF with formatting preservation
Compatible with Linux environments
"""

import sys
import json
import traceback
from pathlib import Path
import tempfile
import os

def convert_word_to_pdf_professional(input_path, output_path):
    """
    Convert Word document to PDF using python-docx and reportlab
    """
    try:
        from docx import Document
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        import io
        from PIL import Image as PILImage
        
        # Load Word document
        doc = Document(input_path)
        
        # Create PDF document
        pdf_doc = SimpleDocTemplate(
            output_path,
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )
        
        # Get styles and create custom styles
        styles = getSampleStyleSheet()
        
        # Create custom styles for better formatting
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Title'],
            fontSize=18,
            spaceAfter=12,
            textColor=colors.black,
            alignment=TA_CENTER
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading1'],
            fontSize=14,
            spaceAfter=8,
            spaceBefore=12,
            textColor=colors.black,
            alignment=TA_LEFT
        )
        
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=6,
            textColor=colors.black,
            alignment=TA_JUSTIFY,
            leading=14
        )
        
        # Build content list
        content = []
        
        # Process document paragraphs
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                content.append(Spacer(1, 6))
                continue
            
            # Determine style based on content
            style = normal_style
            
            # Check if it's a title or heading
            if len(text) < 50 and (
                paragraph.style.name.startswith('Heading') or
                text.isupper() or
                (len(text.split()) <= 5 and any(word.istitle() for word in text.split()))
            ):
                if len(text) < 30:
                    style = heading_style
                else:
                    style = title_style
            
            # Handle text formatting (bold, italic)
            formatted_text = escape_html(text)
            
            # Check for bold formatting in runs
            bold_text = False
            for run in paragraph.runs:
                if run.bold and run.text.strip():
                    bold_text = True
                    break
            
            if bold_text:
                formatted_text = f"<b>{formatted_text}</b>"
            
            # Add paragraph to content
            try:
                para = Paragraph(formatted_text, style)
                content.append(para)
                content.append(Spacer(1, 3))
            except Exception as e:
                # Fallback for problematic text
                para = Paragraph(escape_html(text), normal_style)
                content.append(para)
        
        # Process tables
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    row_data.append(cell_text if cell_text else " ")
                table_data.append(row_data)
            
            if table_data:
                # Create table with styling
                pdf_table = Table(table_data)
                pdf_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                    ('FONTSIZE', (0, 1), (-1, -1), 9),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                
                content.append(Spacer(1, 12))
                content.append(pdf_table)
                content.append(Spacer(1, 12))
        
        # Add footer with conversion info
        content.append(Spacer(1, 20))
        footer_style = ParagraphStyle(
            'Footer',
            parent=styles['Normal'],
            fontSize=8,
            textColor=colors.grey,
            alignment=TA_CENTER
        )
        content.append(Paragraph("Converted by I Love Making PDF", footer_style))
        
        # Build PDF
        pdf_doc.build(content)
        
        print(json.dumps({
            "success": True,
            "message": f"Word document successfully converted to PDF with {len(doc.paragraphs)} paragraphs and {len(doc.tables)} tables."
        }))
        return True
        
    except Exception as e:
        print(f"Professional conversion failed: {e}", file=sys.stderr)
        print(traceback.format_exc(), file=sys.stderr)
        return False

def convert_with_mammoth_fallback(input_path, output_path):
    """
    Fallback conversion using mammoth to extract HTML then convert to PDF
    """
    try:
        import mammoth
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import SimpleDocTemplate, Paragraph
        from html import unescape
        import re
        
        # Extract HTML from Word document
        with open(input_path, "rb") as docx_file:
            result = mammoth.convert_to_html(docx_file)
            html_content = result.value
        
        # Clean HTML and extract text
        text_content = re.sub('<[^<]+?>', '\n', html_content)
        text_content = unescape(text_content)
        
        # Create PDF
        pdf_doc = SimpleDocTemplate(output_path, pagesize=A4)
        styles = getSampleStyleSheet()
        content = []
        
        # Split into paragraphs and process
        paragraphs = text_content.split('\n\n')
        for para_text in paragraphs:
            para_text = para_text.strip()
            if para_text:
                para = Paragraph(escape_html(para_text), styles['Normal'])
                content.append(para)
        
        pdf_doc.build(content)
        
        print(json.dumps({
            "success": True,
            "message": "Word document converted to PDF using fallback method"
        }))
        return True
        
    except Exception as e:
        print(f"Fallback conversion failed: {e}", file=sys.stderr)
        return False

def escape_html(text):
    """
    Escape HTML special characters for ReportLab
    """
    if not text:
        return ""
    
    # Replace problematic characters
    text = text.replace('&', '&amp;')
    text = text.replace('<', '&lt;')
    text = text.replace('>', '&gt;')
    text = text.replace('"', '&quot;')
    text = text.replace("'", '&#x27;')
    
    # Handle special characters that might cause issues
    text = text.replace('\x00', '')  # Remove null characters
    text = text.replace('\r\n', '<br/>')
    text = text.replace('\n', '<br/>')
    text = text.replace('\r', '<br/>')
    
    return text

def main():
    if len(sys.argv) != 3:
        print(json.dumps({
            "success": False,
            "error": "Usage: python word_to_pdf_converter.py <input.docx> <output.pdf>"
        }))
        sys.exit(1)
    
    input_path = sys.argv[1]
    output_path = sys.argv[2]
    
    # Validate input file
    if not os.path.exists(input_path):
        print(json.dumps({
            "success": False,
            "error": "Input file not found"
        }))
        sys.exit(1)
    
    if not input_path.lower().endswith('.docx'):
        print(json.dumps({
            "success": False,
            "error": "Input file must be a .docx file"
        }))
        sys.exit(1)
    
    try:
        # Try professional conversion first
        if convert_word_to_pdf_professional(input_path, output_path):
            sys.exit(0)
        
        # Try fallback method
        if convert_with_mammoth_fallback(input_path, output_path):
            sys.exit(0)
        
        # If all methods fail
        print(json.dumps({
            "success": False,
            "error": "All conversion methods failed. Please ensure the Word document is not corrupted."
        }))
        sys.exit(1)
        
    except Exception as e:
        print(json.dumps({
            "success": False,
            "error": f"Conversion failed: {str(e)}"
        }))
        sys.exit(1)

if __name__ == "__main__":
    main()